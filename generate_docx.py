from docx import Document

doc = Document()

doc.add_heading('RoadWatch AI - Project Documentation', 0)

# Topic
doc.add_heading('1. Selected Topic', level=1)
doc.add_paragraph('The selected topic for this hackathon submission is: Road SOS.')

# Project Overview
doc.add_heading('2. Project Overview', level=1)
doc.add_paragraph(
    'RoadWatch AI is a public audit platform and real-time anomaly detection system. '
    'It aims to solve the lack of transparency in municipal road repair budgets and empower '
    'citizens to report, track, and verify road hazards using simulated dashcam feeds.'
)

# Key Features
doc.add_heading('3. Key Features', level=1)
doc.add_paragraph('• Live Map Dashboard: Displays real-time road anomalies (Potholes, Cracks, Faded Markings) '
                  'with glowing markers coded by severity.')
doc.add_paragraph('• Neural Analysis & Dashcam Simulation: Simulates the processing of video feeds to automatically '
                  'detect road surface issues and dispatch actionable telemetry.')
doc.add_paragraph('• Citizen Rewards & Gamification: Users earn experience points (XP) and rank badges for '
                  'auditing routes and verifying hazards, incentivizing public governance.')
doc.add_paragraph('• Road Analytics: Provides comprehensive indices including Ride Comfort, Surface Quality, '
                  'and severity breakdowns across different municipal wards.')

# Tech Stack
doc.add_heading('4. Technologies Used', level=1)
p = doc.add_paragraph()
p.add_run('Frontend: ').bold = True
p.add_run('React.js, Vite, Tailwind CSS (for Glassmorphism UI), Framer Motion, Lucide React.\n')
p.add_run('Mapping: ').bold = True
p.add_run('React-Leaflet, CartoDB Dark Base Maps.\n')
p.add_run('Backend & API: ').bold = True
p.add_run('Python, FastAPI, Uvicorn.')

# Instructions to Run
doc.add_heading('5. Instructions to Run the Project', level=1)
doc.add_paragraph('The project is divided into two parts: the frontend and the backend. Both must be running concurrently.')

doc.add_heading('Starting the Backend:', level=2)
doc.add_paragraph('1. Open a terminal and navigate to the backend directory: e:\\ROAD SAFETY PROJECT\\backend')
doc.add_paragraph('2. Install dependencies (if not already installed): pip install fastapi uvicorn')
doc.add_paragraph('3. Start the server: python main.py')
doc.add_paragraph('The backend API will run on http://127.0.0.1:8000.')

doc.add_heading('Starting the Frontend:', level=2)
doc.add_paragraph('1. Open a new terminal and navigate to the frontend directory: e:\\ROAD SAFETY PROJECT\\frontend')
doc.add_paragraph('2. Install node modules: npm install')
doc.add_paragraph('3. Start the development server: npm run dev')
doc.add_paragraph('The web application will be accessible at http://localhost:5173.')

# How to use
doc.add_heading('6. How to Use the App', level=1)
doc.add_paragraph('• Once the frontend is loaded, navigate to the "Control Center" in the sidebar and click "Live Dashcam Feed".')
doc.add_paragraph('• This triggers the backend AI simulation which generates anomalies on the map.')
doc.add_paragraph('• Click on any glowing marker on the map to instantly view the Full Report for that specific anomaly.')
doc.add_paragraph('• Use the sidebar navigation to switch between the Live Map, Citizen Rewards, and Road Analytics views.')

doc.save('RoadWatch_AI_Documentation.docx')
print('Documentation saved successfully as RoadWatch_AI_Documentation.docx')
